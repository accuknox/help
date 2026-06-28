# -*- coding: utf-8 -*-
"""Build the RBI Model Risk Management white paper as a branded DOCX, starting
from utils/doc-ppt-template/WORD_TEMPLATE_ACCUKNOX.docx so it keeps the
template's embedded Space Grotesk font, heading styles, and brand blues.

Reads as a final copy: an executive summary up top, an "RBI asks / AccuKnox
delivers" map, inline links instead of paragraph numbers, and readable fonts."""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPO = r"D:\Atharva\AccuKnox\HelpDocs"
TEMPLATE = os.path.join(REPO, "utils", "doc-ppt-template", "WORD_TEMPLATE_ACCUKNOX.docx")
OUT = os.path.join(REPO, "utils", "rbi-mrm-2026",
                   "AccuKnox-RBI-Model-Risk-Management-Whitepaper.docx")
IMG = os.path.join(REPO, "utils", "rbi-mrm-2026", "whitepaper-images")
LOGO = os.path.join(REPO, "docs", "assets", "images", "logo-black.png")

BLUE = RGBColor(0x00, 0x3B, 0xF6)
NAVY = RGBColor(0x0D, 0x1B, 0x4B)
GREY = RGBColor(0x59, 0x59, 0x59)
HDR_FILL = "E3EAFD"
ALT_FILL = "F4F6FE"
RT_HYPERLINK = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

# Links reused from the blog, placed inline at the right spots.
L_AIML = "https://help.accuknox.com/how-to/aiml-overview/"
L_SHADOW = "https://accuknox.com/blog/shadow-ai-security-explained"
L_REDTEAM = "https://help.accuknox.com/use-cases/red-teaming/"
L_MODELARMOR = "https://help.accuknox.com/use-cases/modelarmor/"
L_FIREWALL = "https://help.accuknox.com/use-cases/prompt-firewall-overview/"
L_FIREWALL_BLOG = "https://accuknox.com/blog/stateful-prompt-firewall-guardrail-for-ai-security"
L_RUNTIME = "https://accuknox.com/blog/zero-trust-runtime-security-for-ai-age"
L_AIDR = "https://help.accuknox.com/use-cases/aidr/"
L_RBI = "https://www.rbi.org.in/scripts/BS_PressReleaseDisplay.aspx?prid=63006"
L_RBI_SEBI = "https://accuknox.com/blog/rbi-and-sebi-compliance"
L_RBI_SBOM = "https://accuknox.com/blog/rbi-sbom-mandate-banking-compliance-platform"

doc = Document(TEMPLATE)

body = doc.element.body
for el in list(body):
    if el.tag == qn("w:sectPr"):
        continue
    body.remove(el)


# ---------- helpers ----------
def _shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def para(text="", size=11, bold=False, italic=False, color=NAVY,
         align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
        if color is not None:
            r.font.color.rgb = color
    return p


def bullet(text, size=11):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(4); pf.left_indent = Inches(0.3); pf.first_line_indent = Inches(-0.18)
    r = p.add_run("•\t"); r.font.size = Pt(size); r.font.color.rgb = BLUE
    r2 = p.add_run(text); r2.font.size = Pt(size); r2.font.color.rgb = NAVY
    return p


def numbered(n, text, size=11):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(4); pf.left_indent = Inches(0.3); pf.first_line_indent = Inches(-0.3)
    r = p.add_run("%d.\t" % n); r.font.size = Pt(size); r.font.bold = True; r.font.color.rgb = BLUE
    r2 = p.add_run(text); r2.font.size = Pt(size); r2.font.color.rgb = NAVY
    return p


def h1(text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(8)
    return p


def h2(text):
    p = doc.add_paragraph(text, style="Heading 2")
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    return p


def add_hyperlink(paragraph, text, url, size=11):
    r_id = paragraph.part.relate_to(url, RT_HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "003BF6"); rPr.append(col)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text; run.append(t)
    h.append(run); paragraph._p.append(h)


def mapping(asks, delivers, stays=None, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
    r = p.add_run("RBI asks "); r.font.bold = True; r.font.size = Pt(size); r.font.color.rgb = BLUE
    r2 = p.add_run(asks); r2.font.size = Pt(size); r2.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4 if stays else 10)
    r3 = p2.add_run("AccuKnox delivers "); r3.font.bold = True; r3.font.size = Pt(size); r3.font.color.rgb = BLUE
    for seg in delivers:
        if isinstance(seg, tuple):
            add_hyperlink(p2, seg[0], seg[1], size=size)
        else:
            rr = p2.add_run(seg); rr.font.size = Pt(size); rr.font.color.rgb = NAVY

    if stays:
        p3 = doc.add_paragraph(); p3.paragraph_format.space_after = Pt(10)
        r4 = p3.add_run("Stays with you "); r4.font.bold = True; r4.font.size = Pt(size); r4.font.color.rgb = GREY
        r5 = p3.add_run(stays); r5.font.size = Pt(size); r5.font.color.rgb = NAVY


def rule(color="003BF6", size=12):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(size)); b.set(qn("w:space"), "1"); b.set(qn("w:color"), color)
    pBdr.append(b)
    succ = {"shd", "tabs", "spacing", "ind", "jc", "rPr"}
    anchor = next((c for c in pPr if c.tag.split("}")[-1] in succ), None)
    (anchor.addprevious if anchor is not None else pPr.append)(pBdr)
    return p


def page_break():
    p = doc.add_paragraph(); p.add_run().add_break()
    br = p.runs[0]._r.find(qn("w:br"))
    if br is not None:
        br.set(qn("w:type"), "page")


def add_toc():
    p = doc.add_paragraph(); run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); f1.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click and choose Update Field, or press Ctrl+A then F9, to build the contents."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for e in (f1, instr, f2, t, f3):
        run._r.append(e)


def _cell_margins(table, top=70, bottom=70, left=110, right=110):
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        e = OxmlElement("w:" + edge); e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa"); mar.append(e)
    succ = {"tblLook", "tblCaption", "tblDescription"}
    anchor = next((c for c in tblPr if c.tag.split("}")[-1] in succ), None)
    (anchor.addprevious if anchor is not None else tblPr.append)(mar)


def _borders(table, color="C4CCDE", sz=4):
    tblPr = table._tbl.tblPr
    bs = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz)); e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        bs.append(e)
    succ = {"tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription"}
    anchor = next((c for c in tblPr if c.tag.split("}")[-1] in succ), None)
    (anchor.addprevious if anchor is not None else tblPr.append)(bs)


def add_table(headers, rows, widths, header_size, body_size):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        _borders(table)
    table.autofit = False
    _cell_margins(table)
    # repeat the header row when the table breaks across pages
    trPr = table.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)
    hc = table.rows[0].cells
    for i, htext in enumerate(headers):
        hc[i].width = Inches(widths[i]); _shade(hc[i], HDR_FILL)
        cp = hc[i].paragraphs[0]; cp.paragraph_format.space_after = Pt(2); cp.paragraph_format.space_before = Pt(2)
        r = cp.add_run(htext); r.font.bold = True; r.font.size = Pt(header_size); r.font.color.rgb = NAVY
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].width = Inches(widths[ci])
            if ri % 2 == 1:
                _shade(cells[ci], ALT_FILL)
            cp = cells[ci].paragraphs[0]; cp.paragraph_format.space_after = Pt(3); cp.paragraph_format.space_before = Pt(3)
            r = cp.add_run(val); r.font.size = Pt(body_size); r.font.color.rgb = NAVY
            if ci == len(row) - 1:
                r.font.bold = True
                if val.startswith("Direct"):
                    r.font.color.rgb = BLUE
    return table


def figure(img_name, number, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(os.path.join(IMG, img_name), width=Inches(4.8))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run("Figure %d. " % number); r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = BLUE
    r2 = cap.add_run(caption); r2.font.size = Pt(9.5); r2.font.italic = True; r2.font.color.rgb = GREY


# ======================= FRONT COVER =======================
para(space_after=2); para(space_after=2); para(space_after=2)
para("FOR RBI-REGULATED ENTITIES", size=11, bold=True, color=BLUE, space_after=10)
title = doc.add_paragraph(style="Title")
title.add_run("Meeting the RBI Guidance on Model Risk Management").font.color.rgb = NAVY
sub = doc.add_paragraph(); sub.paragraph_format.space_before = Pt(4)
sr = sub.add_run("A security control map for AI and ML models"); sr.font.size = Pt(16); sr.font.color.rgb = BLUE; sr.font.bold = True
rule()
para("This paper takes the Reserve Bank of India's Guidance on Regulatory "
     "Principles for Model Risk Management and maps each AI and ML requirement to "
     "the security controls AccuKnox provides, and to the duties that stay with "
     "the regulated entity.", size=11.5, space_after=10)
para("Applies to commercial and co-operative banks, NBFCs, payments and small "
     "finance banks, all-India financial institutions such as NABARD and EXIM "
     "Bank, asset reconstruction companies, and credit information companies.",
     size=11, color=GREY, space_after=20)
if os.path.exists(LOGO):
    lp = doc.add_paragraph(); lp.paragraph_format.space_before = Pt(40)
    lp.add_run().add_picture(LOGO, width=Inches(2.0))
para("AccuKnox   ·   White Paper   ·   Final   ·   June 2026", size=11, bold=True, color=NAVY, space_before=6)
page_break()

# ======================= CONTENTS =======================
h1("Contents")
add_toc()
page_break()

# ======================= 1. EXECUTIVE SUMMARY =======================
h1("1. Executive summary")
es = doc.add_paragraph(); es.paragraph_format.space_after = Pt(8)
es.add_run("The Reserve Bank of India has set out how regulated entities must "
           "govern the risk of the models they run, including models that use AI "
           "and machine learning, in its ").font.size = Pt(11)
add_hyperlink(es, "Guidance on Regulatory Principles for Model Risk Management", L_RBI)
r = es.add_run(". It reaches commercial and co-operative banks, NBFCs, payments "
               "and small finance banks, all-India institutions, and credit "
               "information companies, and is open for comment until 24 July 2026.")
r.font.size = Pt(11); r.font.color.rgb = NAVY
para("Most of the guidance is technology-neutral. The AI and ML chapter is where "
     "the bulk of the security work sits, and it is where most institutions have "
     "the least in place today.", size=11)
para("This paper maps each AI and ML requirement to a control you can deploy and "
     "the evidence it produces, and it is honest about the parts that stay with "
     "you. AccuKnox is not a governance framework. It supplies the inventory, the "
     "validation results, the runtime controls, and the audit trail your Model "
     "Risk Management Framework relies on.", size=11)
para("At a glance, AccuKnox covers:", size=11, bold=True, space_after=4, space_before=2)
for b in [
    "A live inventory of every AI model, dataset, and pipeline, including shadow assets.",
    "Independent, repeatable validation of model behaviour through automated red teaming.",
    "Runtime controls for customer-facing and generative models: a stateful prompt "
    "firewall, operating-system-level enforcement, and a reachable kill-switch.",
    "Continuous monitoring that catches drift, behaviour change, and silent provider updates.",
    "Documented, traceable evidence for your board, your committee, and your auditor.",
]:
    bullet(b)
para("What stays with you: the framework itself, risk-tier decisions, model "
     "soundness and fairness, customer disclosures, and contractual terms.",
     size=11, space_before=4)
page_break()

# ======================= 2. ABOUT =======================
h1("2. About this paper")
para("This paper concentrates on the AI and ML parts of the RBI guidance, because "
     "that is where most of the security work sits, along with the deployment and "
     "monitoring duties that go with them. It is written for the security and risk "
     "teams who will operate these controls.")
para("One boundary keeps it honest. AccuKnox is a security platform, not a "
     "governance framework. It does not replace your Model Risk Management "
     "Framework. It produces the inventory, the test results, the runtime "
     "controls, and the audit trail that the framework relies on. Where a "
     "requirement is yours to own and not ours to solve, this paper says so.")
h2("Who this applies to")
para("RBI names eleven categories of regulated entity. In practice it reaches "
     "almost everyone the RBI supervises:", space_after=4)
for b in [
    "Commercial banks, including foreign banks, and small finance, payments, "
    "local area, and regional rural banks",
    "Urban and rural co-operative banks",
    "NBFCs across the base, middle, upper, and top layers",
    "All-India financial institutions: EXIM Bank, NABARD, NaBFID, NHB, and SIDBI",
    "Asset reconstruction companies and credit information companies",
]:
    bullet(b)
para("If a model has a material effect on a business decision, RBI expects you to "
     "govern it, and that includes models that use AI and ML.", space_before=4)

# ======================= 3. WHO OWNS WHAT =======================
h1("3. Who owns what")
para("Model risk management is shared work. The regulated entity owns the "
     "framework and the judgement calls. AccuKnox owns the technical controls and "
     "the evidence. Reading the map with that split in mind avoids the common "
     "mistake of buying a tool and assuming the obligation is met.")
add_table(
    ["The regulated entity owns", "Where AccuKnox helps"],
    [
        ["Board-approved framework, risk appetite, approval committees",
         "Inventory, validation results, and monitoring reports the board and committee review"],
        ["The risk-tier decision for each model",
         "Risk signals that inform tiering: exposure, autonomy, customer reach"],
        ["Conceptual and statistical soundness, fairness statistics",
         "Behavioural and security validation through automated red teaming"],
        ["Training-data governance and quality",
         "Runtime data protection and behaviour-change signals"],
        ["Telling users they deal with AI, and human-handoff design",
         "Enforcement: block, kill-switch, and surfacing outputs for human review"],
        ["Contractual audit rights and the business continuity policy",
         "Audit trail, continuous monitoring, and traceable findings"],
    ],
    widths=[3.15, 3.15], header_size=11, body_size=10.5)
page_break()

# ======================= 4. REQUIREMENT MAP =======================
h1("4. The requirement map")
para("Each item states what RBI asks for, then what AccuKnox delivers against it, "
     "then any part that stays with you.", color=GREY, space_after=6)

h2("4.1 Governance and accountability")
mapping("for a board-approved framework covering every model, with the board "
        "setting risk appetite, a board committee overseeing it, and senior "
        "management running it. You stay accountable for every model outcome, "
        "in-house or vendor-supplied.",
        ["the evidence that framework runs on, not the framework itself: a current "
         "inventory of AI assets, validation results from red teaming, and "
         "monitoring reports your committee can review."],
        "the governance structure, the committees, and the risk appetite.")

h2("4.2 A complete inventory of every AI model")
mapping("for an accurate, current inventory of all models, active, inactive, and "
        "retired, with upstream and downstream dependencies, and enhanced "
        "documentation for AI models so they can be traced and audited.",
        ["automatic ", ("discovery of your AI estate", L_AIML),
         ": models, datasets, compute, and the pipelines that connect them, across "
         "cloud and on-prem. It surfaces the ", ("shadow models", L_SHADOW),
         " nobody registered, and keeps the inventory current because discovery "
         "runs continuously. Findings are tagged to recognised AI risk frameworks "
         "for traceability."],
        "spreadsheet-style models in the wider definition belong in your GRC register.")

h2("4.3 Risk-based model tiering")
mapping("you to classify every model by how material and how complex it is, "
        "review the tier at least once a year, and weigh the autonomy and reliance "
        "placed on an AI model's output.",
        ["the signals that make the tiering decision evidence-based: which models "
         "face customers, how exposed they are, what they can reach, and what red "
         "teaming found."],
        "the tier decision itself.")

h2("4.4 Independent validation and structured challenge")
mapping("for independent validation of every model, before and after deployment "
        "and on every material change, and structured challenge, red-teaming above "
        "all, for anything customer-facing or generative, tested under stressed "
        "and adversarial inputs.",
        [("automated red teaming", L_REDTEAM),
         " against your models for prompt injection, jailbreaks, hallucination, "
         "toxic output, and unsafe code, before launch and again on every update, "
         "with each run documented as evidence."],
        "conceptual and statistical soundness, which is your validation team's work.")

h2("4.5 Third-party and hosted models")
mapping("you to stay accountable for a vendor model whatever assurance it carries, "
        "validate it yourself, and account for supply-chain risk and "
        "provider-driven change. Where a provider discloses little, limit usage.",
        ["red teaming against hosted models such as AWS Bedrock, NVIDIA Triton, and "
         "vLLM on your terms, and ", ("model-artifact scanning", L_MODELARMOR),
         " for tampering such as pickle-deserialization payloads and poisoned "
         "weights. When disclosure is thin, the prompt firewall enforces the usage "
         "limits RBI asks for."],
        "contractual audit rights and vendor documentation, which you negotiate.")

h2("4.6 Explainability and compensating controls")
mapping("you to set explainability thresholds and, where full explainability is "
        "not achievable, to apply compensating controls: corroborate output before "
        "use, validate more often, monitor continuously, and restrict usage.",
        ["exactly that layer. You set the threshold; when a model falls short "
         "AccuKnox supplies the controls, response checks that verify output before "
         "a customer sees it, scheduled re-tests, continuous monitoring, and hard "
         "usage limits at the prompt boundary."])

h2("4.7 Hallucination, bias, and consumer protection")
mapping("for control boundaries against hallucination in generative models, "
        "identification of bias and discriminatory output, and a flat rule against "
        "deploying a model that harms consumers.",
        ["hallucination and toxicity measurement during red teaming, and a ",
         ("stateful prompt firewall", L_FIREWALL),
         " that inspects responses in production, blocking toxic content, leaked "
         "PII, and off-policy answers before the customer sees them."],
        "statistical fairness testing on protected groups.")

h2("4.8 Deployment controls")
mapping("that deploying a model not introduce vulnerabilities, naming access "
        "controls, safeguards against cyber risk, and the risks from external "
        "interfaces, APIs, and integration pipelines.",
        ["least-privilege enforcement ", ("at the operating-system layer", L_RUNTIME),
         ", so a model talked into misbehaving still cannot act. A runtime policy "
         "blocks an attempt to read a credentials file whatever the prompt says, "
         "and API security covers the interface and pipeline risks."])

h2("4.9 Customer-facing and generative models")
mapping("for models that interact with customers: defenses against prompt "
        "injection and adversarial input, limits on session and context "
        "persistence, detection of anomalous usage, a clear disclosure that the "
        "user is dealing with AI, and a route to human help.",
        ["a ", ("stateful prompt firewall", L_FIREWALL_BLOG),
         " that scores the whole conversation, the thing that catches multi-turn "
         "jailbreaks a single-prompt filter misses. It blocks injection and "
         "adversarial input, caps tokens and context, and AI-DR flags anomalous "
         "usage."],
        "the AI disclosure and the human-handoff route, which are choices in your application.")

h2("4.10 Human oversight")
mapping("for human oversight of AI models, including override, suspension, and "
        "kill-switch arrangements, and periodic human review of model-driven "
        "decisions.",
        ["the enforcement side. Runtime protection is the kill-switch: it blocks or "
         "deactivates model behaviour at the operating-system layer, independent of "
         "the model's own guardrails, and flagged outputs surface for human review."],
        "who reviews, how often, and how reviewers are trained against over-reliance.")

h2("4.11 Monitoring, change, and automatic updates")
mapping("for ongoing monitoring of every deployed model, attention to data and "
        "concept drift, and tighter controls for models that update automatically.",
        [("AI-DR", L_AIDR),
         " that watches models in production, flags behaviour change after a "
         "provider-driven or automatic update, and triggers a fresh red-team run so "
         "a quiet update does not bypass validation."],
        "statistical drift on your training data, which sits in your MLOps pipeline.")

h2("4.12 Business continuity and decommissioning")
mapping("you to plan for model failure with fallback mechanisms, tell stakeholders "
        "when a model is retired, and keep decommissioned models in the inventory "
        "for at least ten years.",
        ["an inventory that tracks active, inactive, and retired AI assets, and "
         "runtime enforcement that can act as a fallback by denying unsafe "
         "behaviour if a model degrades."],
        "the continuity policy and the ten-year retention record.")

h2("4.13 Selection, development, and data governance")
mapping("you to define the rationale before building a model, follow a structured "
        "development process, govern the data, and make sure models are not "
        "overfitted and do not rely on spurious correlations.",
        ["a narrower contribution here: scanning model artifacts for tampering. The "
         "rest is model-quality work."],
        "this work, owned by your data and modelling teams. AccuKnox does not "
        "assess overfitting or statistical soundness.")
page_break()

# ======================= 5. CONSOLIDATED MAP =======================
h1("5. Consolidated map")
para("Coverage key. Direct: AccuKnox provides the control. Supports: AccuKnox "
     "supplies evidence or signals that feed a process you own. Yours: you own "
     "this and AccuKnox does not provide it.", size=10.5, color=GREY, space_after=8)
rows = [
    ["Board framework, committee oversight, accountability", "Reports and evidence for the board and committee", "Supports"],
    ["Inventory of all models and dependencies, AI documentation", "Continuous AI asset discovery, dependency view, framework tagging", "Direct"],
    ["Risk-based model tiering", "Exposure, autonomy, and customer-reach signals", "Supports"],
    ["Do not deploy models that harm consumers", "Output filtering of toxic, PII, and off-policy responses", "Supports"],
    ["Selection, development, data governance, model quality", "Artifact scanning only", "Yours"],
    ["Independent validation, before and after, periodic", "Automated red teaming, documented and repeatable", "Direct"],
    ["Approval and exception structure", "Validation evidence for approval", "Supports"],
    ["Deployment and ongoing monitoring", "Continuous runtime and behaviour monitoring (AI-DR)", "Direct"],
    ["Change management and material-change re-validation", "Behaviour-change detection, re-run red teaming, findings log", "Supports"],
    ["Business continuity, decommissioning, retention", "Inventory of inactive assets, fallback enforcement", "Supports"],
    ["Third-party accountability and independent validation", "Red teaming of hosted models, artifact scans", "Direct"],
    ["AI scope, limited-disclosure mitigants, supply chain", "Usage limits, artifact scans, provider-update detection", "Direct"],
    ["Explainability thresholds and compensating controls", "Monitoring, output corroboration, usage limits", "Direct"],
    ["Hallucination control boundaries", "Red teaming and response checks", "Direct"],
    ["Bias and discriminatory output", "Toxicity testing and output filtering", "Supports"],
    ["Data risk, data and concept drift", "Runtime data protection, behaviour-change signal", "Supports"],
    ["Structured challenge and red-teaming", "Automated red teaming", "Direct"],
    ["Controls for automatic updates", "Re-validate on update, monitor behaviour change", "Direct"],
    ["Deployment controls: access, cyber, API and pipeline", "Runtime and zero-trust enforcement, API security", "Direct"],
    ["Prompt injection, adversarial input, session limits, anomaly detection", "Stateful prompt firewall and AI-DR", "Direct"],
    ["AI disclosure and human-handoff option", "Application UX", "Yours"],
    ["Human oversight, kill-switch, periodic review", "Runtime kill and override, surfaced outputs", "Direct"],
]
add_table(["Requirement", "AccuKnox capability", "Coverage"],
          rows, widths=[2.5, 3.0, 0.95], header_size=10.5, body_size=9.5)
page_break()

# ======================= 6. DEPLOYMENT =======================
h1("6. Deploying in a regulated environment")
para("Most of the duties above involve sending model inputs and outputs through a "
     "control plane. For an RBI-regulated entity, where that control plane runs, "
     "and where the data goes, matters as much as the controls themselves.")
para("AccuKnox runs as SaaS or fully on-premises, including air-gapped. In an "
     "on-premises deployment the prompt firewall, red teaming, discovery, and "
     "runtime enforcement run inside your own environment, so prompts, responses, "
     "and model telemetry do not leave the bank's boundary. That keeps the "
     "controls compatible with data-localisation expectations and with an internal "
     "audit that needs to inspect the system directly.")
para("Findings, alerts, and audit logs export to the SIEM and ticketing tools you "
     "already run, so model risk evidence lands in the same place as the rest of "
     "your security and audit trail.")

# ======================= 7. WHERE TO START =======================
h1("7. Where to start")
para("You do not have to do everything at once. A practical order that follows "
     "RBI's own risk logic:", space_after=4)
for i, b in enumerate([
    "Discover the inventory. You cannot tier or validate what you cannot see, and "
    "shadow models are where the surprises are.",
    "Red team the customer-facing and generative models first. RBI singles them "
    "out, and they carry the most consumer risk.",
    "Put the stateful prompt firewall in front of those endpoints to cover prompt "
    "injection, session limits, and harmful output.",
    "Turn on runtime enforcement and continuous monitoring so a bypassed guardrail "
    "or a silent model update does not become an incident.",
], 1):
    numbered(i, b)
para("Each step produces evidence your Model Risk Management Framework can record, "
     "which is what turns a security deployment into a compliance position.",
     space_before=4)
page_break()

# ======================= SOURCES / BACK COVER =======================
h1("Sources and further reading")
h2("The RBI guidance")
src = doc.add_paragraph(); src.paragraph_format.space_after = Pt(10)
src.add_run("Reserve Bank of India, ").font.size = Pt(11)
add_hyperlink(src, "Guidance on Regulatory Principles for Model Risk Management", L_RBI)
r = src.add_run(", issued 24 June 2026, open for public comment until 24 July 2026.")
r.font.size = Pt(11); r.font.color.rgb = NAVY

h2("AccuKnox documentation")
for label, url in [
    ("AI and ML security onboarding", L_AIML),
    ("Prompt Firewall", L_FIREWALL),
    ("AI Red Teaming", L_REDTEAM),
    ("AI Detection and Response (AI-DR)", L_AIDR),
    ("Model artifact scanning (ModelArmor)", L_MODELARMOR),
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    rr = p.add_run(label + ":  "); rr.font.size = Pt(11); rr.font.bold = True; rr.font.color.rgb = NAVY
    add_hyperlink(p, url, url, size=10)

h2("Related reading")
for label, url in [
    ("AccuKnox for RBI and SEBI compliance", L_RBI_SEBI),
    ("The RBI SBOM mandate and banking compliance", L_RBI_SBOM),
    ("Stateful prompt firewall: a guardrail for AI security", L_FIREWALL_BLOG),
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    rr = p.add_run(label + ":  "); rr.font.size = Pt(11); rr.font.bold = True; rr.font.color.rgb = NAVY
    add_hyperlink(p, url, url, size=10)

rule()
para("AccuKnox is a zero trust CNAPP that secures cloud, workloads, and AI systems "
     "from build to runtime. This paper maps RBI's Model Risk Management guidance "
     "to security controls and is provided for information, not as legal or "
     "compliance advice.", size=10, color=GREY, space_before=8)
page_break()

# ======================= APPENDIX =======================
h1("Appendix A. Screenshots")
para("The diagrams and screenshots below illustrate the controls referenced in "
     "this paper. The first four are architecture and flow diagrams; the rest are "
     "from the AccuKnox console.", size=10.5, color=GREY, space_after=8)
figs = [
    ("diagram-ai-spm-architecture.png", "The AI-SPM security architecture: "
     "discovery across managed and unmanaged infrastructure, a deployment pipeline "
     "with scanning, and the prompt firewall in front of cloud LLMs."),
    ("managed-onprem-deployments.png", "The AI footprint AccuKnox discovers, across "
     "managed cloud services and on-prem stacks."),
    ("prompt-firewall-pipeline.png", "The stateful prompt firewall inspection "
     "pipeline: normalize, classify, contextualize, score, enforce."),
    ("ai-dr-workflow.png", "The AI-DR workflow, from event logs and telemetry "
     "through detection to automated remediation."),
    ("fig01-console-overview.png", "A single console for cloud, workload, and AI "
     "findings, the place model risk evidence is collected."),
    ("fig02-ai-inventory-models.png", "AI application inventory with per-model "
     "detail, the inventory RBI asks for."),
    ("fig03-pipeline-topology.png", "AI pipeline topology showing the user region, "
     "model, and backend, the dependency picture the inventory must capture."),
    ("fig04-red-teaming.png", "Automated red teaming results across prompt "
     "injection, hallucination, code generation, and sentiment, the structured "
     "challenge RBI asks for."),
    ("fig05-prompt-firewall-policy.png", "Prompt firewall policy applied to the "
     "prompt side, the response side, or both, for the customer-facing controls."),
    ("fig06-response-code-block.png", "A response policy that strips code from "
     "model output, an example of an output control boundary."),
    ("fig07-runtime-block-aws.png", "Application behaviour monitoring shows a "
     "blocked access to the .aws directory, runtime enforcement in action."),
    ("fig08-runtime-policy.png", "A runtime protection policy that denies the "
     "unsafe action at the operating-system layer, the kill-switch in practice."),
    ("fig09-zero-trust-discovery.png", "Zero trust policy discovery derives the "
     "least-privilege access a workload actually needs, supporting access controls."),
    ("fig10-findings.png", "Findings grouped for triage and ticketing, the ongoing "
     "monitoring evidence RBI expects."),
    ("fig11-rules-engine.png", "A rules engine for custom control rules, for "
     "example flagging any asset missing runtime protection."),
    ("fig12-compliance-frameworks.png", "Findings mapped to control frameworks, the "
     "traceable and auditable evidence RBI expects."),
]
for i, (name, cap) in enumerate(figs, 1):
    figure(name, i, cap)

# ---- sectPr last + gutter ----
sectPr = body.find(qn("w:sectPr"))
if sectPr is not None:
    body.remove(sectPr); body.append(sectPr)
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is not None and pgMar.get(qn("w:gutter")) is None:
        pgMar.set(qn("w:gutter"), "0")

# ---- header: replace template placeholder, keep its logo ----
sec = doc.sections[0]
sec.different_first_page_header_footer = True
for hp in sec.header.paragraphs:
    for hr in hp.runs:
        if "Document Title" in (hr.text or ""):
            hr.text = "RBI Model Risk Management: A Security Control Map"
            hr.font.size = Pt(9); hr.font.color.rgb = GREY

# ---- footer: page number only ----
footer = sec.footer
fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Page "); fr.font.size = Pt(8); fr.font.color.rgb = GREY
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
fr2 = OxmlElement("w:r"); rpr = OxmlElement("w:rPr"); sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16"); rpr.append(sz)
fr2.append(rpr); t = OxmlElement("w:t"); t.text = "1"; fr2.append(t); fld.append(fr2); fp._p.append(fld)

# ---- update fields (TOC) on open, correct schema position ----
settings = doc.settings.element
if settings.find(qn("w:updateFields")) is None:
    uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true")
    after = {"hdrShapeDefaults", "footnotePr", "endnotePr", "compat", "rsids",
             "mathPr", "attachedSchema", "themeFontLang", "clrSchemeMapping",
             "doNotIncludeSubdocsInStats", "decimalSymbol", "listSeparator",
             "docId", "chartTrackingRefBased"}
    anchor = next((c for c in settings if c.tag.split("}")[-1] in after), None)
    (anchor.addprevious if anchor is not None else settings.append)(uf)

doc.save(OUT)
print("Saved:", OUT)
