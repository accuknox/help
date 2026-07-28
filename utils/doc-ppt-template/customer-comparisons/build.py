# -*- coding: utf-8 -*-
"""Build the AccuKnox-branded SonarQube Cloud comparison doc from the master Word template."""
import copy
from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r"D:\Atharva\AccuKnox\HelpDocs\utils\doc-ppt-template\customer-comparisons\AccuKnox_vs_SonarQubeCloud_Comparison.docx"

NAVY = RGBColor(0x11, 0x20, 0x6D)
BLUE = RGBColor(0x00, 0x3B, 0xF6)
BODY = RGBColor(0x0D, 0x1B, 0x4B)
MUTED = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEADER_FILL = "11206D"
ACCUKNOX_FILL = "EEF3FF"
ALT_FILL = "F5F6FF"
BORDER_COLOR = "C4CCDE"

TITLE = "AccuKnox ASPM vs. SonarQube Cloud"
SUBTITLE = "Customer Comparison"

INTRO = (
    "SonarQube Cloud (formerly SonarCloud) is a well established code quality and SAST "
    "platform, and many engineering teams already rely on it inside the SDLC. AccuKnox ASPM "
    "is built to extend that same rigor beyond the pull request, fusing SAST, DAST, SCA, "
    "secrets, IaC, and SBOM with API security, AI/LLM security, CNAPP, and runtime enforcement "
    "in a single, correlated risk graph. The two platforms are frequently run together: "
    "SonarQube Cloud strengthens code quality and developer feedback, while AccuKnox closes "
    "the gap from code to cloud to runtime."
)

ROWS = [
    ("AI-SAST",
     ["Tiered AI-accelerated SAST: deterministic rules, fine-tuned LLM triage, and large-model review in one pipeline",
      "Findings prioritized using reachability and runtime context, so teams focus on what's actually exploitable"],
     ["AI Code Assurance and Agentic Analysis identify and gate AI-generated code",
      "AI CodeFix offers AI assisted remediation suggestions for supported rule violations"]),
    ("Rule Engine & Automation",
     ["Automated ticketing, ownership routing, and noise reduction across all domains"],
     ["No automated ticketing or routing engine",
      "Custom rule authoring for teams standardizing code-quality policy"]),
    ("Quality Gates",
     ["Policy-as-code gates that can block a release based on code, container, IaC, or API risk, not code alone",
      "Gates can factor in live exploitability, not just static severity"],
     ["Go/no-go quality gate at the pull-request level",
      "Branch analysis and PR decoration with configurable pass/fail thresholds"]),
    ("IDE Extension",
     ["Findings and fix guidance surfaced in-editor through IDE plugin across VS Code, Cursor, IntelliJ, and Windsurf"],
     ["SonarQube for IDE plugin across VS Code, IntelliJ, Cursor, and Windsurf"]),
    ("SBOM & Dependency Graph",
     ["xBOM generates SBOM, CBOM, QBOM, and AI-BOM, covering software, cryptography, and AI components",
      "Dependency graph mapped to attack paths and runtime reachability, aligned to CERT-In and RBI reporting needs"],
     ["No native SBOM generation",
      "SCA identifies vulnerable open-source dependencies within the code-scanning workflow"]),
    ("SCA",
     ["Integrated SCA across proprietary and open-source dependencies"],
     ["Available as part of SonarQube Advanced Security, reviewing AI, first-party, and open-source code"]),
    ("DAST",
     ["Native DAST in the same pipeline",
      "Authenticated TOTP-based scans supported",
      "Coverage for OWASP Top 10 with Baseline and Extended Scan"],
     ["Not a native capability; typically paired with a separate DAST tool"]),
    ("API Security",
     ["Dedicated module for shadow API discovery, inventory, and attack-surface scanning",
      "API findings correlated with runtime and gateway telemetry"],
     ["Not a native capability; API risk is addressed at the code level only"]),
    ("AI / LLM Security",
     ["AI-SPM covers prompt guardrails, AI red-teaming, AI detection and response, and agentic/MCP security across 100+ models",
      "Extends protection to deployed AI applications and agents, not just AI-authored code"],
     ["Not part of the product; no AI/LLM security capabilities"]),
    ("CNAPP (CSPM / CIEM)",
     ["Full Zero Trust CNAPP, CSPM, CWPP, KSPM, CDR, CIEM, and DSPM unified with ASPM in one risk graph"],
     ["Not part of the product; SonarQube Cloud is scoped to code analysis"]),
    ("CWPP and Runtime Security",
     ["eBPF/LSM-based runtime enforcement (KubeArmor, a CNCF Sandbox project) with minimal performance overhead"],
     ["Not part of the product; SonarQube Cloud is scoped to code analysis"]),
    ("Runtime Security",
     ["Inline runtime prevention that blocks anomalous and zero-day behavior, not just static detection",
      "Runtime telemetry feeds back into prioritization to reduce alert volume"],
     ["Not part of the product; analysis is complete at build/PR time"]),
    ("Deployment Model",
     ["SaaS, on-premises, or fully air-gapped; suited to regulated and classified environments"],
     ["SaaS only; on-prem/air-gapped deployment requires the separate SonarQube Server product"]),
    ("Integration Support",
     ["ITSM support via rules engine automation",
      "Can be integrated with a SIEM tool to route findings",
      "Dynamic notification flows can be created using the Rules Engine",
      "Flexibility to integrate other tools like SonarQube and CheckMarx"],
     ["Limited integration support, present for ITSM and notifications"]),
    ("Reports",
     ["Executive and detailed reports in different formats can be extracted",
      "Scheduled and on-demand reports"],
     ["Limited reporting capabilities"]),
]

TAKEAWAYS = [
    "A single risk graph from code to cloud to runtime, reducing the manual correlation work required when SAST, DAST, API, and cloud findings live in separate tools.",
    "Native DAST and API security testing, so dynamic and API risk don't require a second platform and a separate remediation workflow.",
    "Runtime-verified prioritization via eBPF/LSM enforcement (KubeArmor, a CNCF Sandbox project), so teams act on what's exploitable in production rather than every static finding.",
    "A dedicated AI-SPM layer covering prompt security, AI red-teaming, and agentic/MCP protection, extending coverage to deployed AI systems, not just AI-authored code.",
    "Deeper supply-chain visibility through xBOM (SBOM, CBOM, QBOM, AI-BOM), supporting CERT-In, RBI, and similar regulatory reporting.",
    "Flexible deployment, SaaS, on-premises, or fully air-gapped, for teams with data-residency or classified-environment requirements.",
]


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color=BORDER_COLOR, sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def set_col_widths(table, widths_in):
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Emu(int(widths_in[idx] * 914400))
    grid = table._tbl.find(qn("w:tblGrid"))
    for idx, gridcol in enumerate(grid.findall(qn("w:gridCol"))):
        gridcol.set(qn("w:w"), str(int(widths_in[idx] * 1440)))


def style_run(run, size=9, bold=False, color=BODY, font="Space Grotesk"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font)


def add_bullet_paragraph(cell, text, first, size=9, color=BODY, bold_lead=None):
    p = cell.add_paragraph() if not first else cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.first_line_indent = Pt(-10)
    bullet_run = p.add_run("•  ")
    style_run(bullet_run, size=size, bold=True, color=BLUE, font="Space Grotesk")
    text_run = p.add_run(text)
    style_run(text_run, size=size, bold=False, color=color, font="Space Grotesk")
    return p


def add_plain_paragraph(cell, text, first, size=9, bold=True, color=NAVY):
    p = cell.add_paragraph() if not first else cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color, font="Space Grotesk")
    return p


def main():
    doc = Document(DOC_PATH)

    # --- Landscape orientation with tighter margins for the wide comparison table ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_w, new_h = section.page_height, section.page_width
    section.page_width, section.page_height = new_w, new_h
    section.left_margin = section.right_margin = Emu(int(0.6 * 914400))
    section.top_margin = section.bottom_margin = Emu(int(0.5 * 914400))
    section.header_distance = Emu(int(0.3 * 914400))
    section.footer_distance = Emu(int(0.3 * 914400))
    page_width_in = new_w / 914400

    # --- Header: replace placeholder title text, re-anchor logo to the right margin ---
    header = section.header
    hp = header.paragraphs[0]
    hp.runs[0].text = TITLE
    style_run(hp.runs[0], size=10, bold=True, color=MUTED, font="Space Grotesk")
    anchor = hp.runs[1]._element.find(qn("w:drawing") if False else qn("w:drawing"))
    # locate the anchor drawing in the header run and switch to right-margin relative positioning
    for r in hp.runs:
        drawing = r._element.find(qn("w:drawing"))
        if drawing is not None:
            anchor_el = drawing.find(qn("wp:anchor"))
            posH = anchor_el.find(qn("wp:positionH"))
            for child in list(posH):
                posH.remove(child)
            posH.set("relativeFrom", "rightMargin")
            align = OxmlElement("wp:align")
            align.text = "right"
            posH.append(align)

    # --- Clear placeholder body paragraphs ---
    body = doc.element.body
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # --- Title + subtitle + intro ---
    title_p = doc.add_paragraph(style="Heading 1")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(TITLE)
    style_run(r, size=26, bold=True, color=BLUE, font="Space Grotesk")

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    r = sub_p.add_run(SUBTITLE.upper())
    style_run(r, size=11, bold=True, color=MUTED, font="Space Grotesk")

    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(16)
    intro_p.paragraph_format.line_spacing = 1.15
    r = intro_p.add_run(INTRO)
    style_run(r, size=10, bold=False, color=BODY, font="Space Grotesk")

    # --- Comparison table ---
    n_rows = len(ROWS) + 1
    table = doc.add_table(rows=n_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cap_w = page_width_in * 0.16
    rest = (page_width_in - cap_w) / 2
    set_col_widths(table, [cap_w, rest, rest])

    headers = ["Capability", "AccuKnox ASPM", "SonarQube Cloud"]
    for c, text in enumerate(headers):
        cell = table.cell(0, c)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_borders(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        style_run(r, size=10.5, bold=True, color=WHITE, font="Space Grotesk")

    for i, (cap, ak_points, sq_points) in enumerate(ROWS):
        row_idx = i + 1
        alt = ALT_FILL if i % 2 == 1 else "FFFFFF"

        cap_cell = table.cell(row_idx, 0)
        set_cell_shading(cap_cell, alt)
        set_cell_borders(cap_cell)
        set_cell_margins(cap_cell)
        add_plain_paragraph(cap_cell, cap, first=True, size=9.5, bold=True, color=NAVY)

        ak_cell = table.cell(row_idx, 1)
        set_cell_shading(ak_cell, ACCUKNOX_FILL)
        set_cell_borders(ak_cell)
        set_cell_margins(ak_cell)
        for j, pt in enumerate(ak_points):
            add_bullet_paragraph(ak_cell, pt, first=(j == 0))

        sq_cell = table.cell(row_idx, 2)
        set_cell_shading(sq_cell, alt)
        set_cell_borders(sq_cell)
        set_cell_margins(sq_cell)
        for j, pt in enumerate(sq_points):
            add_bullet_paragraph(sq_cell, pt, first=(j == 0), color=MUTED)

    # keep header row from splitting across pages
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)

    # --- Takeaways section (kept together on its own page) ---
    takeaway_title = doc.add_paragraph(style="Heading 2")
    takeaway_title.paragraph_format.page_break_before = True
    r = takeaway_title.add_run("What This Means for Your Team")
    style_run(r, size=16, bold=True, color=BLUE, font="Space Grotesk")
    takeaway_title.paragraph_format.space_before = Pt(0)
    takeaway_title.paragraph_format.space_after = Pt(8)

    for t in TAKEAWAYS:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.first_line_indent = Pt(-14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.keep_with_next = True
        b = p.add_run("•  ")
        style_run(b, size=10, bold=True, color=BLUE, font="Space Grotesk")
        r = p.add_run(t)
        style_run(r, size=10, bold=False, color=BODY, font="Space Grotesk")

    doc.save(DOC_PATH)
    print("Saved:", DOC_PATH)


if __name__ == "__main__":
    main()
